import sys
import subprocess
import os

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def add_tbl_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def apply_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="1B365D"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_document():
    doc = Document()
    
    # Page setup - Standard 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base colors
    COLOR_PRIMARY = RGBColor(27, 54, 93)   # Deep Navy #1B365D
    COLOR_SECONDARY = RGBColor(74, 85, 104) # Cool Gray #4A5568
    COLOR_TEXT = RGBColor(45, 55, 72)      # Dark Charcoal #2D3748

    # Style setups
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # Document Header / Cover style Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("DOKUMEN PERANCANGAN SISTEM: STOXAREA")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    run_sub = p_sub.add_run("Sistem Pendukung Keputusan Rekomendasi Saham berbasis Machine Learning (XGBoost, SHAP, dan SAW)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY

    # Add Divider Line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_div = p_div.add_run("―" * 40)
    run_div.font.color.rgb = COLOR_SECONDARY
    p_div.paragraph_format.space_after = Pt(24)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.bold = True
        p.add_run(text)
        return p

    # --- 1. PROSES BISNIS ---
    add_heading_1("1. PROSES BISNIS (BUSINESS PROCESS & SYSTEM WORKFLOW)")
    
    p = doc.add_paragraph()
    p.add_run("Ekosistem sistem STOXAREA bekerja dalam alur kerja terintegrasi yang menjembatani data pasar saham dinamis (real-time market data), kecerdasan buatan (Machine Learning), sistem pendukung keputusan (SPK) berjenjang tiga lapis (Three-Tier SPK), dan simulasi transaksi (Virtual Trading). Alur proses bisnis sistem dijabarkan secara detail sebagai berikut:")

    add_heading_2("A. Pengumpulan Data & Preprocessing (Ingestion Layer)")
    p = doc.add_paragraph()
    p.add_run("Tahap Pengumpulan Data dalam sistem merupakan fondasi krusial yang memastikan model AI bekerja dengan data yang bersih, relevan, dan stabil. Alur kerja pada tahap ini dirancang sebagai berikut:")
    
    add_bullet("Sistem menggunakan Automated Pipeline untuk meminimalkan intervensi manual dan memastikan keberlanjutan data. Komponen Ingestor Data bertugas melakukan request ke API Yahoo Finance untuk mengambil data mentah secara terjadwal. Komponen Scheduler mengatur waktu penarikan data secara berkala (setiap penutupan pasar pukul 17:00 WIB) agar data di dashboard selalu mutakhir.", "Arsitektur Automated Pipeline: ")
    add_bullet("Sistem terhubung secara mendalam dengan API Yahoo Finance untuk menarik dua kategori data utama: (1) Data Riwayat Harga (OHLCV) mencakup harga Pembukaan (Open), Tertinggi (High), Terendah (Low), Penutupan (Close), dan Volume transaksi sebagai bahan baku utama model XGBoost untuk mencari pola teknikal. (2) Metrik Rasio Keuangan menarik data fundamental seperti ROE, DER, dan PBV (diimplementasikan menggantikan PER untuk stabilitas data emiten ber-laba negatif) secara otomatis untuk memberikan konteks kesehatan bisnis pada setiap emiten.", "Integrasi API Yahoo Finance & Struktur Data: ")
    add_bullet("Untuk menjaga agar server dan perangkat keras tidak mengalami overhead akibat beban kerja komputasi, sistem menerapkan pemisahan jalur data: (1) Bulk Historical Ingestion menarik data riwayat dalam jumlah besar (5 tahun terakhir) secara terpisah saat beban server rendah (malam hari) karena mengonsumsi memori dan bandwidth besar. (2) Dynamic Ticker Update melakukan penarikan data harga terkini secara ringan dan dinamis untuk memperbarui nilai portofolio atau ticker di dasbor secara real-time demi responsivitas aplikasi.", "Strategi Pemisahan Arus Data (Server Stability): ")
    add_bullet("Dari total 900+ emiten di BEI, tidak semuanya layak untuk dianalisis oleh AI karena adanya saham 'tidur' atau tidak likuid. Sistem menyaring emiten berdasarkan volume transaksi harian rata-rata 30 hari > 10.000 lot, hari aktif bursa setahun >= 100 hari, dan harga penutupan terakhir > Rp 50. Hasil akhirnya adalah sekitar 500 emiten saham teraktif di BEI yang siap dianalisis, menjamin XGBoost hanya belajar dari pergerakan harga wajar.", "Filtrasi Emiten Aktif (Filtering Stage): ")
    add_bullet("Setelah data disaring dan dibersihkan, data fundamental dan teknikal disimpan secara terpusat ke database (PostgreSQL backend) dalam format terstruktur yang siap diproses oleh tahap Feature Engineering.", "Penyimpanan Terpusat: ")

    add_heading_2("B. Personalisasi Pengguna (SPK Lapis 1 - User Profiling)")
    p = doc.add_paragraph()
    p.add_run("Proses bisnis yang berinteraksi langsung dengan pengguna dimulai dari tahap orientasi (onboarding) melalui pengisian kuesioner psikologis Likert 1-5 berjumlah 10 pertanyaan (masing-masing kriteria diukur oleh 2 pertanyaan) untuk mengategorikan tipe risiko investasi pengguna:")
    
    add_bullet("Kuesioner dirancang mengukur lima dimensi: K1 (Target Keuntungan - matched dengan AI Score XGBoost), K2 (Kualitas Perusahaan - matched dengan ROE), K3 (Toleransi Risiko - matched dengan DER), K4 (Sensitivitas Harga - matched dengan PBV), dan K5 (Kapasitas Finansial - matched dengan Logika Veto/Safety Net).", "Kriteria Profiling: ")
    add_bullet("Skor jawaban Likert bernilai 1 (Konservatif/Aman), 3 (Moderat/Seimbang), atau 5 (Agresif/Volatil). Total skor dari K1 hingga K5 akan menentukan kategori profil investasi pengguna: Skor < 12 diklasifikasikan sebagai Konservatif (mengutamakan keamanan modal dan dividen stabil), skor 12 s.d. 18 sebagai Moderat (mencari keseimbangan pertumbuhan dan keamanan), dan skor > 18 sebagai Agresif (fokus penuh pada capital gain dan momentum).", "Skoring dan Klasifikasi Profil: ")
    add_bullet("Setelah profil terbentuk, sistem menetapkan bobot preferensi (total wajib 1.0) untuk rumus SAW di SPK Lapis 3: (1) Profil Konservatif menetapkan bobot AI Score = 0.10, ROE = 0.45, DER = 0.35, PBV = 0.10. (2) Profil Moderat menetapkan bobot AI Score = 0.35, ROE = 0.30, DER = 0.15, PBV = 0.20. (3) Profil Agresif menetapkan bobot AI Score = 0.60, ROE = 0.10, DER = 0.10, PBV = 0.20.", "Matriks Bobot Preferensi: ")
    add_bullet("Kriteria ke-5 (Kapasitas Finansial) berfungsi sebagai Guardrail etis/sabuk pengaman melalui Override Logic. Jika pengguna menjawab dana investasi berasal dari 'Dana Darurat' atau 'Uang Pokok' (K5 = 1), atau K5 = 3 (dana sedang) tetapi K3 (toleransi risiko) = 1 (sangat takut rugi), sistem secara otomatis melakukan Veto dan menurunkan profil risiko ke tingkat paling aman (Konservatif) terlepas dari seberapa tinggi skor pada kriteria lainnya. Ini menghentikan anomali ambisi agresif di atas kerentanan finansial nyata.", "Fitur Veto (K5 Safety Net): ")

    add_heading_2("C. Penyaringan Data Saham (SPK Lapis 2 - Stock Scoring)")
    p = doc.add_paragraph()
    p.add_run("Tahap ini bekerja secara sepenuhnya objektif untuk menilai performa emiten di bursa efek tanpa melibatkan emosi atau preferensi subjektif pengguna:")
    
    add_bullet("Sistem menyaring saham sleeping stocks (saham gocap) dan mengklasifikasikan emiten qualified ke dalam 11 sektor resmi Bursa Efek Indonesia (BEI) ditambah 1 sektor pertanian untuk penyajian informasi yang rapi.", "Pra-pemrosesan & Klasifikasi Sektoral: ")
    add_bullet("Sistem menggunakan algoritma Machine Learning XGBoost untuk membaca pola kompleks secara simultan dari perpaduan data fundamental dan 11 indikator teknikal (seperti RSI, MACD, MA). Hasilnya berupa probabilitas kenaikan harga saham dalam 5 hari ke depan yang disebut AI Score.", "Pemrosesan AI (XGBoost): ")
    add_bullet("Emiten dievaluasi berdasarkan 4 kriteria utama: (1) K1: AI Score (Benefit - makin tinggi makin baik, mewakili momentum bursa), (2) K2: ROE (Benefit - makin tinggi makin baik, mewakili efisiensi manajemen mencetak laba), (3) K3: DER (Cost - makin rendah makin bagus, mewakili stabilitas keuangan dari utang), dan (4) K4: PBV (Cost - makin rendah makin bagus, mewakili kewajiban valuasi harga murah/undervalued).", "Dimensi Kriteria Penilaian Objektif: ")
    add_bullet("Nilai mentah dikonversi menjadi skala desimal seragam antara 0 hingga 1 menggunakan rumus normalisasi Min-Max, menghasilkan Matriks Rating R berdimensi N x 4 (misalnya: [0.71, 0.73, 1.00, 0.33] untuk BBCA). Matriks data objektif inilah yang akan dikalikan dengan bobot kepentingan pengguna di SPK Lapis 3.", "Normalisasi dan Pembentukan Matriks Rating: ")
    add_bullet("Sebagai finalisasi analisis objektif, sistem melakukan pemotongan data (slicing) untuk mengambil 10 hingga 15 saham teratas di setiap sektor. Saham dalam top-15 ini otomatis adalah emiten paling unggul secara teknis dibandingkan rekan se-sektornya dan siap disajikan di halaman Market bursa.", "Seleksi Top-N (Slicing & Kualitas): ")

    add_heading_2("D. Pemrosesan Machine Learning (XGBoost & SHAP)")
    p = doc.add_paragraph()
    p.add_run("XGBoost berperan sebagai mesin prediksi utama yang mengubah data pasar mentah menjadi sinyal probabilitas terukur di SPK Lapis 2 dengan akselerasi CUDA pada GPU Nvidia:")
    
    add_bullet("Sistem menyatukan data teknikal (RSI, MACD, Moving Average untuk mendeteksi momentum dan titik jenuh jenuh beli/jual) dan data fundamental (ROE, DER, PBV untuk mendeteksi jangkar keamanan bisnis) sebagai bahan baku training.", "Sinergi Data Input (Bahan Baku): ")
    add_bullet("Melalui Feature Engineering, data mentah ditransformasikan dengan teknik lagging data untuk menangkap pola tren jangka pendek. Pembelajaran berjenjang (Residual Learning) membangun serangkaian pohon keputusan secara berurutan guna meminimalkan kesalahan prediksi (residual error) secara bertahap dengan optimasi fungsi objektif terregularisasi.", "Alur Kerja XGBoost: ")
    add_bullet("Setelah prediksi probabilitas dihasilkan, model memanfaatkan algoritma SHAP (TreeExplainer) untuk membedah 'kotak hitam' keputusan AI. SHAP mengekstrak 3 kontributor utama dan menerjemahkannya ke dalam bahasa alami (misalnya: 'ROE tinggi mendorong naik probabilitas') di dasbor emiten untuk transparansi retail.", "Transparansi melalui SHAP (Explainable AI): ")
    add_bullet("Probabilitas biner yang dihasilkan kemudian dikalibrasi menggunakan Isotonic Regression (CalibratedClassifierCV) untuk menghasilkan nilai riil probabilitas kenaikan harga [0.0, 1.0]. AI Score dan ROE diperlakukan sebagai benefit, sedangkan DER dan PBV sebagai cost.", "Output & Kalibrasi: ")

    add_heading_2("E. Agregasi & Peringkat Berbasis SAW (SPK Lapis 3)")
    p = doc.add_paragraph()
    p.add_run("SPK Lapis 3 bertindak sebagai otak pengambil keputusan final yang mempertemukan profil psikologis subjektif pengguna (SPK 1) dengan matriks rating objektif pasar (SPK 2):")
    
    add_bullet("Sistem mempertemukan kriteria psikologis dengan variabel pasar: Target Keuntungan dijodohkan dengan AI Score (Ambisi vs Probabilitas), Kualitas Perusahaan dijodohkan dengan ROE (Idealisme vs Efisiensi), Toleransi Risiko dijodohkan dengan DER (Ketahanan Mental vs Solvabilitas), dan Sensitivitas Harga dijodohkan dengan PBV (Gaya Belanja vs Kewajaran Harga).", "Matriks Jembatan Kriteria (Matching Table): ")
    add_bullet("Melalui agregasi SAW, skor kecocokan akhir Vi dihitung: V_i = sum(w_j * r_ij). Nilai dinyatakan dalam persentase (0-100%). Sebelum pemeringkatan final, filter rule-based blacklist dijalankan untuk mengeliminasi saham berbahaya sesuai profil (contoh: DER > 150% langsung dibuang bagi tipe Konservatif). Hal ini menjamin rekomendasi aman dari risiko fatal.", "Logika Agregasi & Filter Blacklist: ")
    add_bullet("Seluruh komputasi SPK 3 dioptimalkan berjalan di atas infrastruktur tangguh dengan sistem Double-Checked Locking Cache per profil (Konservatif, Moderat, Agresif) guna mencegah lag komputasi massal (thundering herd). Hasil akhir matching tertinggi diekstrak menjadi Top Picks berisi 3 hingga 5 saham rekomendasi terbaik untuk keseluruhan bursa dan 1-2 saham per sektor tampilan sektoral.", "Ekosistem Komputasi & Penyajian Bersih: ")

    add_heading_2("F. Fitur Market & Dasbor Publik (Market Analysis Page)")
    p = doc.add_paragraph()
    p.add_run("Menyediakan halaman analisis pasar objektif untuk memantau saham potensial secara umum tanpa terpengaruh profil pengguna:")
    add_bullet("Menampilkan daftar 30 s.d. 150 saham bursa teraktif yang diurutkan murni berdasarkan momentum AI Score tertinggi, dilengkapi sparkline historis 7 hari dan tag sentimen harian otomatis (Bullish/Bearish/Netral). Halaman dilindungi sistem cache dinamis TTL 5 menit untuk mencegah throttling API Yahoo Finance.", "Top AI Momentum & Optimasi Cache: ")

    add_heading_2("G. Virtual Trading (Simulator Transaksi Realistis)")
    p = doc.add_paragraph()
    p.add_run("Simulator transaksi dirancang semirip mungkin dengan mekanisme perdagangan riil di Bursa Efek Indonesia (BEI):")
    add_bullet("Validasi pembelian menggunakan harga pasar terkini ditambah biaya transaksi riil broker Indonesia sebesar 0.15%. Portofolio diupdate dengan metode Average Cost untuk melacak modal beli secara akurat. Penjualan dikenakan biaya 0.25% (mencakup komisi broker, PPN, dan PPh Final 0.1% penjualan). Aksi korporasi seperti Stock Split BBCA 1:5 tervalidasi admin secara otomatis menyesuaikan kuantitas (Qty x 5) dan harga modal (Avg Price / 5).", "Transaksi & Resolusi Aksi Korporasi: ")

    # --- 2. RANCANGAN DATASET ---
    add_heading_1("2. RANCANGAN DATASET (DATASET DESIGN)")
    p = doc.add_paragraph()
    p.add_run("Rancangan struktur dataset yang digunakan dalam sistem STOXAREA terbagi menjadi tiga jenis representasi data: Data Mentah (Raw), Data Terproses (Processed Features), dan Skema Basis Data Relasional Operasional.")

    add_heading_2("A. Dataset OHLCV Saham Historis (Format CSV per Emiten)")
    p = doc.add_paragraph()
    p.add_run("Sumber Data: Yahoo Finance API (yfinance). Lokasi: data/raw/ohlcv/{TICKER}.csv (Contoh: BBCA.JK.csv).")

    # Table 1: Raw OHLCV
    table1 = doc.add_table(rows=8, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table1)
    
    headers1 = ["Nama Atribut", "Tipe Data", "Kunci", "Keterangan / Deskripsi Fitur"]
    col_widths1 = [Inches(1.2), Inches(1.0), Inches(0.8), Inches(3.5)]
    
    # Style Header Row
    hdr_cells = table1.rows[0].cells
    add_tbl_header(table1.rows[0])
    for idx, name in enumerate(headers1):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    data1 = [
        ["Date", "Datetime", "Foreign", "Tanggal transaksi bursa (Format: YYYY-MM-DD)"],
        ["Open", "Float", "-", "Harga pembukaan perdagangan bursa pada hari terkait (IDR)"],
        ["High", "Float", "-", "Harga tertinggi saham pada hari terkait (IDR)"],
        ["Low", "Float", "-", "Harga terendah saham pada hari terkait (IDR)"],
        ["Close", "Float", "-", "Harga penutupan akhir perdagangan bursa (IDR)"],
        ["Volume", "BigInt", "-", "Jumlah lembar saham yang ditransaksikan pada hari terkait"],
        ["ticker", "String", "Foreign", "Kode ticker unik saham IDX bersuffix .JK"]
    ]

    for row_idx, row_data in enumerate(data1, 1):
        row = table1.rows[row_idx]
        add_cant_split(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx in [1, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_idx == 0:
                p.runs[0].font.bold = True
        # Alternating background colors
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F7FAFC")

    # Set Column Widths
    for row in table1.rows:
        for idx, width in enumerate(col_widths1):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_2("B. Dataset Fitur dan Target Terproses (Processed Features & Targets)")
    p = doc.add_paragraph()
    p.add_run("Lokasi: data/processed/features_targets.csv. Berisi 11 fitur hasil rekayasa teknikal untuk training dan inferensi XGBoost.")

    # Table 2: Processed Features
    table2 = doc.add_table(rows=16, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table2)
    
    headers2 = ["Nama Kolom", "Tipe Data", "Batasan Nilai", "Keterangan / Penjelasan Fitur"]
    col_widths2 = [Inches(1.5), Inches(1.0), Inches(1.1), Inches(2.9)]
    
    hdr_cells = table2.rows[0].cells
    add_tbl_header(table2.rows[0])
    for idx, name in enumerate(headers2):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    data2 = [
        ["ticker", "String", "-", "Kode ticker unik saham IDX bersuffix .JK"],
        ["Date", "Datetime", "-", "Tanggal pencatatan baris data"],
        ["log_ret_1d", "Float", "[-0.35, 0.35]", "Log return 1 hari terakhir: ln(Close_t / Close_t-1)"],
        ["log_ret_5d", "Float", "[-0.50, 0.50]", "Log return 5 hari terakhir: ln(Close_t / Close_t-5)"],
        ["ma_20_dist", "Float", "[-1.0, 1.0]", "Jarak harga penutupan ke garis rata-rata pergerakan 20 hari"],
        ["ma_50_dist", "Float", "[-1.5, 1.5]", "Jarak harga penutupan ke garis rata-rata pergerakan 50 hari"],
        ["bb_width", "Float", "[0.0, 1.0]", "Lebar volatilitas Bollinger Bands: (BB_upper - BB_lower) / MA20"],
        ["bb_position", "Float", "[0.0, 1.0]", "Posisi harga penutupan relatif terhadap garis BB (0=Bawah, 1=Atas)"],
        ["rsi_14", "Float", "[0.0, 100.0]", "Indikator Relative Strength Index periode 14"],
        ["macd_norm", "Float", "[-0.1, 0.1]", "Nilai MACD dibagi harga close (normalisasi lintas saham)"],
        ["macd_signal_norm", "Float", "[-0.1, 0.1]", "Nilai MACD Signal dibagi harga penutupan"],
        ["macd_hist_norm", "Float", "[-0.05, 0.05]", "Nilai MACD Histogram dibagi harga penutupan"],
        ["vol_ma_ratio", "Float", "[0.0, 50.0]", "Rasio volume harian bursa terhadap rata-rata volume 20 hari"],
        ["target_5d_up", "Integer", "{0, 1, NaN}", "Label target: 1 jika harga naik dalam 5 hari ke depan, 0 jika tidak"],
        ["is_latest", "Boolean", "{True, False}", "Bernilai True jika baris adalah data penutupan hari ini"]
    ]

    for row_idx, row_data in enumerate(data2, 1):
        row = table2.rows[row_idx]
        add_cant_split(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx in [1, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_idx == 0:
                p.runs[0].font.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F7FAFC")

    for row in table2.rows:
        for idx, width in enumerate(col_widths2):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_2("C. Skema Tabel Database Relasional (PostgreSQL)")
    p = doc.add_paragraph()
    p.add_run("Secara operasional, STOXAREA menggunakan basis data relasional PostgreSQL dengan desain skema tabel sebagai berikut:")

    # Table 3: Database schema users
    add_heading_3("Tabel: users (Identitas & Profil Pengguna)")
    table3 = doc.add_table(rows=9, cols=4)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table3)
    
    col_widths3 = [Inches(1.5), Inches(1.2), Inches(1.8), Inches(2.0)]
    hdr_cells = table3.rows[0].cells
    add_tbl_header(table3.rows[0])
    for idx, name in enumerate(["Nama Kolom", "Tipe Data", "Batasan (Constraints)", "Penjelasan / Keterangan"]):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    data3 = [
        ["id", "Integer", "PRIMARY KEY, AUTO_INCREMENT", "ID unik internal pengguna"],
        ["email", "Varchar(255)", "UNIQUE, NOT NULL", "Alamat surel unik untuk autentikasi login"],
        ["hashed_password", "Varchar(255)", "NOT NULL", "Hash kata sandi terenkripsi aman menggunakan bcrypt"],
        ["full_name", "Varchar(255)", "NOT NULL", "Nama lengkap terdaftar pengguna"],
        ["risk_profile", "Varchar(50)", "NOT NULL", "Klasifikasi profil hasil Tier 1 (konservatif/moderat/agresif)"],
        ["virtual_balance", "Decimal(15,2)", "DEFAULT 100000000.00, NOT NULL", "Saldo tunai simulator transaksi pengguna (IDR)"],
        ["created_at", "Timestamp", "DEFAULT NOW()", "Tanggal pendaftaran akun pertama kali"],
        ["updated_at", "Timestamp", "DEFAULT NOW()", "Tanggal perubahan data/profil terakhir kali"]
    ]

    for row_idx, row_data in enumerate(data3, 1):
        row = table3.rows[row_idx]
        add_cant_split(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx == 0:
                p.runs[0].font.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F7FAFC")

    for row in table3.rows:
        for idx, width in enumerate(col_widths3):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table 4: stocks
    add_heading_3("Tabel: stocks (Data Profil Emiten Saham)")
    table4 = doc.add_table(rows=9, cols=4)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table4)
    
    hdr_cells = table4.rows[0].cells
    add_tbl_header(table4.rows[0])
    for idx, name in enumerate(["Nama Kolom", "Tipe Data", "Batasan (Constraints)", "Penjelasan / Keterangan"]):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    data4 = [
        ["ticker", "Varchar(20)", "PRIMARY KEY", "Kode ticker unik bursa IDX (contoh: BBCA.JK)"],
        ["name", "Varchar(255)", "NOT NULL", "Nama resmi perusahaan tercatat bursa"],
        ["sector", "Varchar(100)", "NOT NULL", "Klasifikasi sektor industri (Bahasa Indonesia)"],
        ["roe", "Decimal(10,2)", "NULLABLE", "Rasio Return on Equity mentah (Persen)"],
        ["der", "Decimal(10,2)", "NULLABLE", "Rasio Debt to Equity mentah (Kali)"],
        ["pbv", "Decimal(10,2)", "NULLABLE", "Rasio Price to Book Value mentah (Kali)"],
        ["is_qualified", "Boolean", "DEFAULT False, INDEXED", "Indikator kelayakan emiten (lolos filter veto Tier 2)"],
        ["updated_at", "Timestamp", "DEFAULT NOW()", "Tanggal pembaruan data fundamental terakhir"]
    ]

    for row_idx, row_data in enumerate(data4, 1):
        row = table4.rows[row_idx]
        add_cant_split(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx == 0:
                p.runs[0].font.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F7FAFC")

    for row in table4.rows:
        for idx, width in enumerate(col_widths3):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table 5: portfolios
    add_heading_3("Tabel: portfolios (Portofolio Kepemilikan Aset)")
    table5 = doc.add_table(rows=8, cols=4)
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table5)
    
    hdr_cells = table5.rows[0].cells
    add_tbl_header(table5.rows[0])
    for idx, name in enumerate(["Nama Kolom", "Tipe Data", "Batasan (Constraints)", "Penjelasan / Keterangan"]):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    data5 = [
        ["id", "Integer", "PRIMARY KEY, AUTO_INCREMENT", "ID unik portofolio aset"],
        ["user_id", "Integer", "FOREIGN KEY (users.id), NOT NULL", "Referensi relasi ke pengguna pemilik portofolio"],
        ["ticker", "Varchar(20)", "NOT NULL, INDEXED", "Kode ticker saham yang dimiliki"],
        ["qty", "Integer", "NOT NULL", "Jumlah kepemilikan unit saham dalam satuan lembar"],
        ["avg_price", "Decimal(15,2)", "NOT NULL", "Rata-rata modal harga beli per lembar saham (IDR)"],
        ["created_at", "Timestamp", "DEFAULT NOW()", "Waktu pembelian perdana saham bursa terkait"],
        ["updated_at", "Timestamp", "DEFAULT NOW()", "Tanggal pembaruan kuantitas saham"]
    ]

    for row_idx, row_data in enumerate(data5, 1):
        row = table5.rows[row_idx]
        add_cant_split(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx == 0:
                p.runs[0].font.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F7FAFC")

    for row in table5.rows:
        for idx, width in enumerate(col_widths3):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table 6: transactions
    add_heading_3("Tabel: transactions (Catatan Transaksi Virtual)")
    table6 = doc.add_table(rows=10, cols=4)
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table6)
    
    hdr_cells = table6.rows[0].cells
    add_tbl_header(table6.rows[0])
    for idx, name in enumerate(["Nama Kolom", "Tipe Data", "Batasan (Constraints)", "Penjelasan / Keterangan"]):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    data6 = [
        ["id", "Integer", "PRIMARY KEY, AUTO_INCREMENT", "ID unik riwayat transaksi virtual"],
        ["user_id", "Integer", "FOREIGN KEY (users.id), NOT NULL", "Relasi ke pengguna yang melakukan transaksi"],
        ["ticker", "Varchar(20)", "NOT NULL, INDEXED", "Kode ticker unik saham yang ditransaksikan"],
        ["type", "Varchar(10)", "NOT NULL", "Jenis transaksi bursa: BUY atau SELL"],
        ["price", "Decimal(15,2)", "NOT NULL", "Harga pelaksanaan per lembar saham pada saat transaksi (IDR)"],
        ["qty", "Integer", "NOT NULL", "Jumlah volume lembar saham yang ditransaksikan"],
        ["fee", "Decimal(15,2)", "NOT NULL, DEFAULT 0.0", "Biaya broker & pajak transaksi riil (0.15% beli, 0.25% jual)"],
        ["net_value", "Decimal(15,2)", "NOT NULL, DEFAULT 0.0", "Nilai bersih akhir kas setelah penyesuaian biaya transaksi"],
        ["timestamp", "Timestamp", "DEFAULT NOW()", "Waktu eksekusi transaksi dicatat oleh sistem"]
    ]

    for row_idx, row_data in enumerate(data6, 1):
        row = table6.rows[row_idx]
        add_cant_split(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx == 0:
                p.runs[0].font.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F7FAFC")

    for row in table6.rows:
        for idx, width in enumerate(col_widths3):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table 7: corporate_action_flags
    add_heading_3("Tabel: corporate_action_flags (Sistem Proteksi & Mitigasi Aksi Korporasi)")
    table7 = doc.add_table(rows=12, cols=4)
    table7.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table7)
    
    hdr_cells = table7.rows[0].cells
    add_tbl_header(table7.rows[0])
    for idx, name in enumerate(["Nama Kolom", "Tipe Data", "Batasan (Constraints)", "Penjelasan / Keterangan"]):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    data7 = [
        ["id", "Integer", "PRIMARY KEY, AUTO_INCREMENT", "ID unik pencatatan bendera aksi korporasi"],
        ["ticker", "Varchar(20)", "NOT NULL, INDEXED", "Kode ticker saham terdeteksi anomali ekstrem"],
        ["prev_close", "Decimal(15,2)", "NULLABLE", "Harga penutupan hari bursa sebelumnya (IDR)"],
        ["curr_close", "Decimal(15,2)", "NULLABLE", "Harga penutupan hari bursa saat ini terdeteksi anomali (IDR)"],
        ["change_pct", "Decimal(5,2)", "NULLABLE", "Persentase perubahan harga ekstrim (>35% dalam sehari)"],
        ["is_resolved", "Boolean", "DEFAULT False, NOT NULL", "Status resolusi manual oleh Administrator (True jika selesai)"],
        ["action_type", "Varchar(50)", "NULLABLE", "Jenis aksi tervalidasi: stock_split, reverse_split, normal_drop"],
        ["split_ratio", "Decimal(10,4)", "NULLABLE", "Rasio pemecahan nilai saham (contoh: 5.0 untuk split 1:5)"],
        ["admin_notes", "Text", "NULLABLE", "Catatan analisis manual oleh Administrator"],
        ["detected_at", "Timestamp", "DEFAULT NOW()", "Waktu sistem mendeteksi lonjakan harga ekstrim harian"],
        ["resolved_at", "Timestamp", "NULLABLE", "Waktu eksekusi aksi penyelarasan portofolio oleh Administrator"]
    ]

    for row_idx, row_data in enumerate(data7, 1):
        row = table7.rows[row_idx]
        add_cant_split(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx == 0:
                p.runs[0].font.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F7FAFC")

    for row in table7.rows:
        for idx, width in enumerate(col_widths3):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table 8: financial_history
    add_heading_3("Tabel: financial_history (Riwayat Fundamental Laporan Keuangan)")
    table8 = doc.add_table(rows=9, cols=4)
    table8.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table8)
    
    hdr_cells = table8.rows[0].cells
    add_tbl_header(table8.rows[0])
    for idx, name in enumerate(["Nama Kolom", "Tipe Data", "Batasan (Constraints)", "Penjelasan / Keterangan"]):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    data8 = [
        ["id", "Integer", "PRIMARY KEY, AUTO_INCREMENT", "ID unik riwayat fundamental perusahaan"],
        ["ticker", "Varchar(20)", "NOT NULL, INDEXED", "Kode ticker saham ekuitas IDX"],
        ["year", "Varchar(4)", "NOT NULL", "Tahun buku pelaporan keuangan (format YYYY, contoh: 2025)"],
        ["revenue", "Decimal(25,2)", "DEFAULT 0, NOT NULL", "Total pendapatan perusahaan dalam tahun berjalan (IDR)"],
        ["net_income", "Decimal(25,2)", "DEFAULT 0, NOT NULL", "Laba bersih setelah pajak perusahaan dalam tahun berjalan (IDR)"],
        ["assets", "Decimal(25,2)", "DEFAULT 0, NOT NULL", "Total aktiva / kekayaan perusahaan (IDR)"],
        ["liabilities", "Decimal(25,2)", "DEFAULT 0, NOT NULL", "Total kewajiban utang jangka pendek & panjang (IDR)"],
        ["equity", "Decimal(25,2)", "DEFAULT 0, NOT NULL", "Total modal ekuitas bersih perusahaan (IDR)"]
    ]

    for row_idx, row_data in enumerate(data8, 1):
        row = table8.rows[row_idx]
        add_cant_split(row)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx == 0:
                p.runs[0].font.bold = True
        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F7FAFC")

    for row in table8.rows:
        for idx, width in enumerate(col_widths3):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # --- 3. METODE YANG DIGUNAKAN ---
    add_heading_1("3. METODE YANG DIGUNAKAN (METHODS USED)")
    
    add_heading_2("A. Algoritma XGBoost Classifier")
    
    p = doc.add_paragraph()
    p.add_run("XGBoost (Extreme Gradient Boosting) dipilih sebagai mesin prediksi utama karena keunggulannya yang superior dalam menangani data tabular terstruktur. Secara akademis, XGBoost merupakan implementasi teroptimasi dari algoritma Gradient Boosted Decision Trees (GBDT). Algoritma ini meminimalkan fungsi kerugian (loss function) melalui penambahan pohon keputusan (decision trees) baru secara berurutan menggunakan teknik optimasi numerik gradient descent.")
    
    p = doc.add_paragraph()
    p.add_run("Secara formal, untuk dataset dengan n sampel dan m fitur, D = {(x_i, y_i)}, prediksi ansambel untuk model aditif dengan K pohon keputusan dirumuskan sebagai:")
    
    p_eq1 = doc.add_paragraph()
    p_eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq1 = p_eq1.add_run("y_pred_i = sum_{k=1}^{K} f_k(x_i),   f_k in F")
    run_eq1.font.italic = True
    run_eq1.font.bold = True
    
    p = doc.add_paragraph()
    p.add_run("Dimana F adalah ruang dari seluruh pohon regresi keputusan (Regression Trees). Untuk meminimalkan kesalahan prediksi sekaligus mencegah overfitting, XGBoost mengimplementasikan fungsi tujuan (objective function) yang ter-regularisasi sebagai berikut:")
    
    p_eq2 = doc.add_paragraph()
    p_eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq2 = p_eq2.add_run("L(phi) = sum_i l(y_i, y_pred_i) + sum_k Omega(f_k)")
    run_eq2.font.italic = True
    run_eq2.font.bold = True

    p = doc.add_paragraph()
    p.add_run("Dengan bagian regularisasi Omega(f) yang mendefinisikan kompleksitas pohon keputusan:")
    
    p_eq3 = doc.add_paragraph()
    p_eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq3 = p_eq3.add_run("Omega(f) = gamma * T + (1/2) * lambda * sum_{j=1}^{T} w_j^2")
    run_eq3.font.italic = True
    run_eq3.font.bold = True

    p = doc.add_paragraph()
    p.add_run("Dimana T adalah jumlah daun (leaves) pada pohon keputusan, w adalah skor bobot pada setiap daun, serta gamma dan lambda adalah parameter regularisasi untuk mengontrol kompleksitas pohon guna mencegah overfitting pada data bursa yang sangat volatil.")

    add_heading_3("Penyelarasan Hiperparameter & Penanganan Imbalance")
    add_bullet("Model dilatih menggunakan konfigurasi parameter optimal: n_estimators=150, max_depth=4, learning_rate=0.05, objective='binary:logistic', eval_metric='logloss', dan scale_pos_weight.", "Konfigurasi Hiperparameter: ")
    add_bullet("Sistem mendeteksi distribusi kelas aktual dan mengatur parameter scale_pos_weight = N_negatif / N_positif secara dinamis. Ini memberikan bobot lebih tinggi pada kelas minoritas (kelas 1 - saham naik >0% dalam 5 hari), sehingga meningkatkan sensitivitas presisi model bursa secara signifikan.", "Imbalance Class Handling: ")
    add_bullet("Agar luaran model mempresentasikan probabilitas dunia nyata secara akurat (bukan sekadar skor peringkat), model di-wrap dengan CalibratedClassifierCV menggunakan metode Isotonic Regression. Hal ini memastikan output AI Score merepresentasikan probabilitas nyata peluang kenaikan harga saham.", "Isotonic Probability Calibration: ")

    add_heading_2("B. Algoritma SHAP (SHapley Additive exPlanations)")
    
    p = doc.add_paragraph()
    p.add_run("Dalam dunia finansial yang sensitif, model kotak hitam (black-box) seperti XGBoost tidak dapat diterima tanpa adanya transparansi keputusan investasi. Oleh karena itu, sistem mengimplementasikan SHAP (SHapley Additive exPlanations).")
    
    p = doc.add_paragraph()
    p.add_run("SHAP didasarkan pada konsep matematis Shapley Values yang berasal dari teori permainan (coalitional game theory). Dalam konteks STOXAREA, fitur-fitur teknikal berperan sebagai pemain, dan perubahan prediksi dibandingkan prediksi rata-rata dasar (Base Value) berperan sebagai hasil permainan (game payout). Kontribusi marjinal rata-rata fitur ke-i dirumuskan secara formal sebagai:")

    p_eq4 = doc.add_paragraph()
    p_eq4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq4 = p_eq4.add_run("phi_i(x) = sum_{S in F \\ {i}} (|S|!(|F| - |S| - 1)! / |F|!) * [ f_x(S U {i}) - f_x(S) ]")
    run_eq4.font.italic = True
    run_eq4.font.bold = True

    p = doc.add_paragraph()
    p.add_run("SHAP value menjamin sifat efisiensi, simetri, dan dummy. Sifat efisiensi menjamin prediksi model dapat diuraikan secara linear sebagai jumlahan dari base value dan kontribusi masing-masing fitur:")
    
    p_eq5 = doc.add_paragraph()
    p_eq5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq5 = p_eq5.add_run("Prediction = Base_Value + sum_{i=1}^{M} phi_i")
    run_eq5.font.italic = True
    run_eq5.font.bold = True

    add_heading_3("Implementasi Teknis SHAP dalam Rekomendasi")
    add_bullet("Sistem memanfaatkan varian TreeExplainer yang sangat efisien untuk model berbasis pohon keputusan guna mengekstrak kontribusi marginal setiap fitur teknikal harian secara instan.", "TreeExplainer Integration: ")
    add_bullet("Sistem mengekstrak 3 fitur dengan kontribusi absolut terbesar. Untuk setiap kontributor utama tersebut, jika bernilai positif diubah menjadi kalimat interpretatif 'mendorong naik probabilitas', dan jika bernilai negatif menjadi 'menekan turun probabilitas'. Ini disajikan dalam bentuk Waterfall Chart interaktif di frontend untuk transparansi analisis bagi investor retail.", "Generasi Insight Bahasa Alami: ")

    # --- METODE SAW EXPLANATION ---
    add_heading_2("C. Algoritma Simple Additive Weighting (SAW)")
    
    p = doc.add_paragraph()
    p.add_run("Metode Simple Additive Weighting (SAW), sering juga dikenal sebagai metode penjumlahan terbobot linear, merupakan salah satu algoritma tertua dan paling banyak digunakan dalam penyelesaian masalah Multi-Attribute Decision Making (MADM). Keunggulan utama dari metode SAW adalah kemampuannya untuk mencari penjumlahan terbobot dari rating kinerja pada setiap alternatif di semua kriteria, menjadikannya sangat intuitif untuk diintegrasikan dengan preferensi subjektif pengguna.")
    
    p = doc.add_paragraph()
    p.add_run("Secara matematis, proses SAW terdiri dari dua langkah utama. Langkah pertama adalah melakukan normalisasi matriks keputusan keputusan (X) ke suatu skala yang dapat diperbandingkan dengan semua rating alternatif yang ada. Rumus normalisasi Min-Max yang digunakan disesuaikan dengan sifat masing-masing kriteria (Benefit atau Cost) sebagai berikut:")
    
    p_eq_saw1 = doc.add_paragraph()
    p_eq_saw1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq_saw1 = p_eq_saw1.add_run("r_ij = (x_ij - min_i(x_ij)) / (max_i(x_ij) - min_i(x_ij))   [untuk Benefit]\n"
                                     "r_ij = 1 - (x_ij - min_i(x_ij)) / (max_i(x_ij) - min_i(x_ij))   [untuk Cost]")
    run_eq_saw1.font.italic = True
    run_eq_saw1.font.bold = True

    p = doc.add_paragraph()
    p.add_run("Langkah kedua adalah melakukan perhitungan nilai preferensi (V_i) untuk setiap alternatif saham dengan cara mengalihkan bobot preferensi (W_j) yang telah diperoleh dari pengisian kuesioner dengan nilai rating ternormalisasi (r_ij):")
    
    p_eq_saw2 = doc.add_paragraph()
    p_eq_saw2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq_saw2 = p_eq_saw2.add_run("V_i = sum_{j=1}^{m} w_j * r_ij")
    run_eq_saw2.font.italic = True
    run_eq_saw2.font.bold = True

    p = doc.add_paragraph()
    p.add_run("Di mana m melambangkan total kriteria (dalam sistem ini berjumlah 4 kriteria) dan w_j adalah bobot preferensi kriteria ke-j. Alternatif dengan nilai V_i yang lebih tinggi menunjukkan tingkat kesesuaian yang lebih optimal terhadap preferensi investasi personal pengguna.")

    # --- METODE SPK EXPLANATION (Three-Tier SPK) ---
    add_heading_2("D. Sistem Pendukung Keputusan Berjenjang (Three-Tier SPK)")
    
    p = doc.add_paragraph()
    p.add_run("STOAXAREA merupakan sebuah Sistem Pendukung Keputusan (SPK) berjenjang tiga lapis (Three-Tier SPK) yang memadukan profil preferensi subjektif pengguna dengan penilaian kuantitatif objektif data pasar. Setiap lapisan dirancang secara formal sebagai berikut:")

    add_heading_3("1. Tier 1 - User Profiling (Profil & Bobot Preferensi Pengguna)")
    p = doc.add_paragraph()
    p.add_run("Berfungsi mengukur preferensi psikologis pengguna dan menerjemahkannya ke dalam parameter kuantitatif vektor bobot preferensi (W) dan label profil risiko.")
    add_bullet("Kuesioner 10 pertanyaan Likert 1-5 mencakup lima dimensi: Ekspektasi Hasil Investasi (K1), Preferensi Kinerja Bisnis (K2), Ketahanan Mental & Risiko (K3), Persepsi Nilai Harga (K4), dan Kapasitas Finansial Nyata (K5). Skor dimensi (K_j) adalah rata-rata kedua pertanyaan per dimensi.", "Input Kuesioner: ")
    p = doc.add_paragraph()
    p.add_run("Vektor bobot preferensi (W) dihitung hanya dari dimensi K1 s.d. K4 menggunakan normalisasi proporsional teoretis:")
    
    p_eq6 = doc.add_paragraph()
    p_eq6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq6 = p_eq6.add_run("w_j = K_j / sum_{i=1}^{4} K_i,  j = 1,2,3,4")
    run_eq6.font.italic = True
    run_eq6.font.bold = True
    
    p = doc.add_paragraph()
    p.add_run("Mekanisme Safety Net (Veto K5): Jika K5 <= 2 (kapasitas finansial sangat rendah), w_1 dikurangi 50% dan selisihnya dipindahkan penuh ke w_3 (toleransi risiko) lalu dinormalisasi ulang agar total sum(W) = 1. Ini mengimplementasikan konsep risk capacity (kemampuan keuangan objektif) yang membedakannya dengan risk tolerance subjektif.")
    p = doc.add_paragraph()
    p.add_run("Penentuan Profil Risiko: Fase pembentukan awal menggunakan K-Means clustering (k=3) dari data minimal 30 pengguna untuk mencari centroid profil: Konservatif, Moderat, dan Agresif. Pada fase operasional harian, pengguna baru dihitung menggunakan Nearest-Centroid (euclidean distance K1-K4 terkecil terhadap centroid profil).")

    add_heading_3("2. Tier 2 - Stock Scoring (Penilaian Objektif Seluruh Saham)")
    p = doc.add_paragraph()
    p.add_run("Melakukan penilaian kuantitatif objektif terhadap seluruh saham qualified di bursa dengan 4 kriteria terpilih (tidak melakukan eliminasi saham):")
    add_bullet("AI Score (Benefit): Probabilitas momentum kenaikan harga jangka pendek horizon 7 hari menggunakan model XGBoost Classifier terkalibrasi.", "Kriteria 1: ")
    add_bullet("ROE (Benefit): Return on Equity mengukur profitabilitas perusahaan dari laporan keuangan.", "Kriteria 2: ")
    add_bullet("DER (Cost): Debt to Equity Ratio mengukur leverage solvabilitas perusahaan.", "Kriteria 3: ")
    add_bullet("PBV (Cost): Price to Book Value mengukur valuasi harga buku saham (dipilih menggantikan PER karena lebih stabil saat laba negatif).", "Kriteria 4: ")
    
    p = doc.add_paragraph()
    p.add_run("Setiap nilai kriteria x_ij saham i dinormalisasi ke rentang [0,1] menggunakan normalisasi Min-Max untuk menyamakan skala kriteria:")
    
    p_eq7 = doc.add_paragraph()
    p_eq7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq7 = p_eq7.add_run("r_ij = (x_ij - min(x_j)) / (max(x_j) - min(x_j))   [untuk Benefit]\n"
                             "r_ij = 1 - (x_ij - min(x_j)) / (max(x_j) - min(x_j))   [untuk Cost]")
    run_eq7.font.italic = True
    run_eq7.font.bold = True
    
    p = doc.add_paragraph()
    p.add_run("Output akhir berupa Matriks Rating R berukuran N x 4 (N = jumlah saham qualified).")

    add_heading_3("3. Tier 3 - Integration, Filtering & Ranking (Integrasi & Pemeringkatan)")
    p = doc.add_paragraph()
    p.add_run("Menggabungkan preferensi personal pengguna (Tier 1) dengan penilaian objektif saham (Tier 2) untuk menghasilkan rekomendasi akhir:")
    
    p = doc.add_paragraph()
    p.add_run("A. Agregasi SAW (Simple Additive Weighting): Skor kecocokan akhir (V_i) dinyatakan dalam persen (0-100%) dengan perkalian matriks berbobot:")
    p_eq8 = doc.add_paragraph()
    p_eq8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq8 = p_eq8.add_run("V_i = sum_{j=1}^{4} w_j * r_ij")
    run_eq8.font.italic = True
    run_eq8.font.bold = True
    
    p = doc.add_paragraph()
    p.add_run("B. Penyaringan Berbasis Aturan (Rule-Based Blacklist): Sebelum pemeringkatan, saham yang tidak memenuhi syarat fundamental disingkirkan. Contoh aturan: untuk profil Konservatif, saham dengan rasio DER mentah > 150% langsung dieliminasi dari daftar rekomendasi akhir.")
    p = doc.add_paragraph()
    p.add_run("C. Pemeringkatan & Top-N: Saham yang lolos penyaringan diurutkan berdasarkan V_i secara menurun. Sebanyak N saham teratas (contoh: Top-5) dipilih sebagai rekomendasi personal terbaik.")

    # --- 4. DIAGRAM RANCANGAN SISTEM & KARDINALITAS ---
    add_heading_1("4. DIAGRAM RANCANGAN SISTEM (SYSTEM DIAGRAMS & CARDINALITIES)")
    
    p = doc.add_paragraph()
    p.add_run("Untuk memvisualisasikan arsitektur sistem secara holistik, berikut adalah spesifikasi alur data, hubungan antarentitas database PostgreSQL, serta struktur pemodelan kelas pemrograman (OOP) yang diimplementasikan pada sistem STOXAREA:")

    add_heading_2("A. Diagram Alir Data Level 0 (DFD Level 0 - Context Diagram)")
    p = doc.add_paragraph()
    p.add_run("DFD Level 0 menggambarkan interaksi sistem secara eksternal dengan entitas luar (Pengguna, Administrator, dan API Yahoo Finance):")
    add_bullet("Pengguna memberikan input berupa Pendaftaran Akun, Kuesioner Profil (K1-K5), dan Perintah Transaksi virtual (BELI/JUAL).", "Input Pengguna: ")
    add_bullet("Sistem mengembalikan hasil berupa Klasifikasi Profil Risiko, Rekomendasi Unggulan (Top Picks SAW), Visualisasi Analisis SHAP, serta data Portofolio & Transaksi.", "Output ke Pengguna: ")
    add_bullet("API Yahoo Finance secara otomatis mengirimkan data mentah riwayat harga saham (OHLCV) dan rasio fundamental (ROE, DER, PBV).", "Aliran API Eksternal: ")
    add_bullet("Administrator melakukan validasi aksi korporasi secara manual dan memicu retrain model ML. Sistem mengirimkan notifikasi deteksi harga ekstrem.", "Aliran Administrator: ")

    add_heading_2("B. Diagram Alir Data Level 1 (DFD Level 1 - Decomposition Diagram)")
    p = doc.add_paragraph()
    p.add_run("DFD Level 1 mendekomposisi proses sistem menjadi 5 modul operasional utama yang berinteraksi aktif dengan database:")
    add_bullet("Memproses pengisian kuesioner pengguna, mendeteksi aturan Veto finansial, menetapkan profil risiko, serta menyimpannya ke tabel users.", "1.0 Penentuan Profil Risiko & Veto: ")
    add_bullet("Secara berkala melakukan ingestion data dari API Yahoo Finance, menyaring saham qualified (is_qualified), serta menyimpan data fundamental ke tabel stocks dan financial_history.", "2.0 Pembersihan & Pengumpulan Data: ")
    add_bullet("Melatih model XGBoost Classifier pada data teknikal, memprediksi AI Score terkalibrasi, menghitung kontribusi fitur harian menggunakan SHAP (TreeExplainer), dan menyimpannya di tabel stocks.", "3.0 Pemodelan Prediksi & Analisis SHAP: ")
    add_bullet("Menghitung normalisasi SAW untuk kriteria benefit (AI Score & ROE) dan kriteria cost (DER & PBV), mengalikannya dengan bobot profil W, melakukan penyaringan blacklist, dan menyajikan Top Picks ke pengguna.", "4.0 Agregasi & Rekomendasi SAW: ")
    add_bullet("Mengeksekusi order virtual BUY/SELL, membebankan fee transaksi (0.15% beli, 0.25% jual), memperbarui data kas dan lot di portfolios & transactions, mendeteksi lonjakan harga ekstrem (>35%), serta memicu penyelarasan portofolio saat Stock Split divalidasi.", "5.0 simulator Transaksi & Aksi Korporasi: ")

    add_heading_2("C. Diagram Kelas UML (UML Class Diagram)")
    p = doc.add_paragraph()
    p.add_run("Diagram kelas UML memetakan struktur pemrograman berbasis objek (OOP models) di backend STOXAREA yang mencakup 6 kelas model relasional utama:")
    add_bullet("Kelas User melacak data identitas pengguna, profil risiko, saldo tunai virtual, serta fungsi pemutakhiran saldo.", "User: ")
    add_bullet("Kelas Stock menyimpan profil dasar emiten saham, nilai fundamental mentah, status kelayakan filter, serta fungsi pengambilan saham qualified.", "Stock: ")
    add_bullet("Kelas Portfolio melacak jumlah kepemilikan lembar saham (qty) dan modal harga beli rata-rata (avg_price) per emiten untuk setiap pengguna.", "Portfolio: ")
    add_bullet("Kelas Transaction menyimpan setiap histori transaksi virtual simulator trading lengkap dengan kuantitas, harga eksekusi, dan komisi sekuritas.", "Transaction: ")
    add_bullet("Kelas CorporateActionFlag mencatat bendera anomali lompatan harga ekstrem harian serta nilai rasio split saham yang tervalidasi.", "CorporateActionFlag: ")
    add_bullet("Kelas FinancialHistory menyimpan data historis laporan keuangan tahunan mencakup pendapatan, laba bersih, aset, utang, dan ekuitas.", "FinancialHistory: ")

    add_heading_2("D. Entity Relationship Diagram (ERD) & Penjelasan Kardinalitas")
    p = doc.add_paragraph()
    p.add_run("Rancangan ERD database PostgreSQL STOXAREA menggunakan relasi 1-to-Many (Satu ke Banyak) yang terikat secara ketat melalui Foreign Key demi menjamin integritas data relasional:")
    add_bullet("1 User dapat memiliki banyak aset saham di portofolionya. Sebaliknya, setiap baris portofolio hanya merujuk pada tepat 1 user.", "USERS ke PORTFOLIOS (1-to-Many): ")
    add_bullet("1 User dapat melakukan banyak aktivitas transaksi beli/jual di simulator. Setiap struk transaksi hanya dilakukan oleh tepat 1 user.", "USERS ke TRANSACTIONS (1-to-Many): ")
    add_bullet("1 Emiten Saham dapat dimiliki di portofolio oleh banyak user yang berbeda. Setiap baris portofolio hanya terikat ke tepat 1 emiten saham.", "STOCKS ke PORTFOLIOS (1-to-Many): ")
    add_bullet("1 Emiten Saham dapat ditransaksikan banyak kali oleh berbagai pengguna. Setiap catatan transaksi terikat ke tepat 1 emiten saham.", "STOCKS ke TRANSACTIONS (1-to-Many): ")
    add_bullet("1 Emiten Saham dapat mengalami banyak bendera anomali harga sepanjang sejarahnya. Setiap baris anomali hanya melaporkan tepat 1 saham.", "STOCKS ke CORPORATE_ACTION_FLAGS (1-to-Many): ")
    add_bullet("1 Emiten Saham memiliki banyak tahun pencatatan laporan keuangan historis. Setiap baris laporan keuangan hanya valid untuk tepat 1 saham.", "STOCKS ke FINANCIAL_HISTORY (1-to-Many): ")
    p = doc.add_paragraph()
    p.add_run("Catatan Hubungan Many-to-Many: Secara logis, terdapat hubungan Banyak-ke-Banyak (Many-to-Many) antara USERS dan STOCKS karena banyak user dapat memiliki banyak saham. Di database PostgreSQL, hubungan ini berhasil dipecah menggunakan tabel perantara PORTFOLIOS dan TRANSACTIONS demi kepatuhan normalisasi database.")

    # Footnote/Metadata
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_foot = p_foot.add_run("Disusun oleh: Kelompok Proyek Sains Data STOXAREA\nUniversitas Putra Bangsa, 2026")
    run_foot.font.italic = True
    run_foot.font.size = Pt(9.5)
    run_foot.font.color.rgb = COLOR_SECONDARY

    # Save document
    output_docx = "c:\\Users\\anmsa\\Downloads\\STOXAREA\\PERANCANGAN_SISTEM_STOXAREA_THREE_TIER.docx"
    try:
        doc.save(output_docx)
        print(f"Document saved successfully at: {output_docx}")
    except PermissionError:
        fallback_docx = "c:\\Users\\anmsa\\Downloads\\STOXAREA\\PERANCANGAN_SISTEM_STOXAREA_THREE_TIER_FINAL.docx"
        doc.save(fallback_docx)
        print(f"Word is locking the original file. Saved updated version to: {fallback_docx}")

if __name__ == "__main__":
    build_document()
